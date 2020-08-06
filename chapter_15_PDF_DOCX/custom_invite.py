#! python3
# custom_invite.py
# Author: Michael Koundouros
"""
Write a program that would generate a Word document with custom invitations. Since Python-Docx can use only those
styles that already exist in the Word document, you will have to first add these styles to a blank Word file and then
open that file with Python-Docx. There should be one invitation per page in the resulting Word document,
so call add_break() to add a page break after the last paragraph of each invitation. This way, you will need to open
only one Word document to print all of the invitations at once.
usage: python custom_invite.py template guestlist
where template is the doxc word file with the invitation template
and guestlist is a txt file with a list of names. One name per line.
"""


import sys
from pathlib import Path
import docx


def invite(template, guestlist):
    doc = docx.Document(template)
    full_text = []
    styles = []
    for para in doc.paragraphs:
        full_text.append(para.text)
        styles.append(para.style)

    with open(guestlist, 'r') as guests:
        guestnames = guests.readlines()

    for guestname in guestnames:
        doc.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
        full_text[1] = guestname.strip('\n')
        for text, style in zip(full_text, styles):
            doc.add_paragraph(text, style=style)

    doc.save('custom_invites.docx')
    return


def main():
    # Make a list of command line arguments, omitting the [0] element
    # which is the script itself.
    args = sys.argv[1:]

    if len(args) != 2:
        print('usage: python custom_invite.py template guestlist')
        sys.exit(1)

    invite(Path(args[0]), args[1])
    print('Done!')


if __name__ == '__main__':
    main()
